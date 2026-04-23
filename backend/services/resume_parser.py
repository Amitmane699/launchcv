"""Rule-based PDF/DOCX resume parser. No AI — regex + heuristics."""
import re
from io import BytesIO
from typing import Dict, Any, List

EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
PHONE_RE = re.compile(r"(?:\+?\d{1,3}[\s-]?)?\(?\d{3,5}\)?[\s.-]?\d{3,4}[\s.-]?\d{4}")
LINKEDIN_RE = re.compile(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+/?", re.I)
WEBSITE_RE = re.compile(r"(?:https?://)?(?:www\.)?[\w-]+\.[a-z]{2,}(?:/\S*)?", re.I)

SECTION_KEYWORDS = {
    "summary":        ["summary", "profile", "objective", "about me"],
    "experience":     ["experience", "work experience", "employment", "professional experience", "work history"],
    "education":      ["education", "academic", "qualifications"],
    "skills":         ["skills", "technical skills", "core competencies", "technologies"],
    "projects":       ["projects", "personal projects", "key projects"],
    "certifications": ["certifications", "certificates", "licenses"],
    "languages":      ["languages"],
    "hobbies":        ["hobbies", "interests"],
}

COMMON_SKILLS = {
    "python", "java", "javascript", "typescript", "react", "angular", "vue", "node",
    "fastapi", "django", "flask", "spring", "express", "nextjs", "nuxt",
    "postgresql", "mysql", "mongodb", "redis", "elasticsearch", "dynamodb", "sql", "nosql",
    "aws", "gcp", "azure", "docker", "kubernetes", "terraform", "ansible", "jenkins",
    "git", "github", "gitlab", "linux", "bash", "rest", "graphql", "grpc",
    "html", "css", "tailwind", "sass", "bootstrap", "figma",
    "tableau", "powerbi", "excel", "sap", "salesforce",
    "photoshop", "illustrator", "indesign", "canva",
    "agile", "scrum", "jira", "confluence",
    "machine learning", "deep learning", "tensorflow", "pytorch", "pandas", "numpy",
    "c++", "c#", ".net", "go", "golang", "rust", "ruby", "rails", "php", "laravel",
    "kotlin", "swift", "android", "ios", "react native", "flutter",
    "marketing", "seo", "sem", "copywriting", "analytics", "ga4",
    "accounting", "finance", "budgeting", "forecasting", "audit",
    "communication", "leadership", "teamwork", "problem solving",
}


def extract_text_from_pdf(data: bytes) -> str:
    import pdfplumber
    text_parts: List[str] = []
    with pdfplumber.open(BytesIO(data)) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            text_parts.append(t)
    return "\n".join(text_parts)


def extract_text_from_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(BytesIO(data))
    parts: List[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    # Also include tables (many resumes use them)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def split_lines(text: str) -> List[str]:
    return [ln.strip() for ln in text.replace("\r", "\n").split("\n") if ln.strip()]


def find_section(lines: List[str], section_keys: List[str]) -> int:
    """Return index of the line that matches one of the section keywords (heading-like)."""
    for i, ln in enumerate(lines):
        low = ln.lower().strip(" :•·")
        if len(low) > 40:
            continue
        for kw in section_keys:
            if low == kw or low.startswith(kw + ":") or low == kw.upper():
                return i
    return -1


def get_section_range(lines: List[str], section_key: str) -> (int, int):
    """Return (start, end) inclusive of a section by key name."""
    start = find_section(lines, SECTION_KEYWORDS.get(section_key, []))
    if start == -1:
        return -1, -1
    # Find next known section after `start`
    end = len(lines)
    for other_key, kws in SECTION_KEYWORDS.items():
        if other_key == section_key:
            continue
        idx = find_section(lines[start + 1:], kws)
        if idx != -1:
            end = min(end, start + 1 + idx)
    return start, end


def parse_personal(text: str, lines: List[str]) -> Dict[str, Any]:
    personal: Dict[str, Any] = {}
    email_m = EMAIL_RE.search(text)
    if email_m:
        personal["email"] = email_m.group(0)
    phone_m = PHONE_RE.search(text)
    if phone_m:
        raw = phone_m.group(0).strip()
        if sum(ch.isdigit() for ch in raw) >= 10:
            personal["phone"] = raw
    li_m = LINKEDIN_RE.search(text)
    if li_m:
        personal["linkedin"] = li_m.group(0)

    # Name: likely the first non-email line near the top
    for ln in lines[:8]:
        if EMAIL_RE.search(ln) or PHONE_RE.search(ln) or "@" in ln:
            continue
        words = ln.split()
        if 2 <= len(words) <= 5 and all(w[0:1].isalpha() for w in words) and not any(w.lower() in {"resume", "curriculum", "cv"} for w in words):
            # Heuristic: name-like if mostly title-cased alpha words
            alpha_titled = sum(1 for w in words if w[0].isupper() and w.replace("-", "").isalpha())
            if alpha_titled >= len(words) - 1:
                personal["fullName"] = " ".join(words)
                break

    # Location heuristic: line with comma, short, near top
    for ln in lines[:12]:
        if "," in ln and len(ln) < 50 and not EMAIL_RE.search(ln) and not PHONE_RE.search(ln) and not personal.get("location"):
            tokens = ln.split(",")
            if all(len(t.strip()) <= 30 for t in tokens) and len(tokens) <= 3:
                personal["location"] = ln.strip()
                break

    return personal


def parse_summary(lines: List[str]) -> str:
    s, e = get_section_range(lines, "summary")
    if s == -1:
        return ""
    body = " ".join(lines[s + 1:e]).strip()
    return body[:600]


def parse_skills(text: str, lines: List[str]) -> List[str]:
    skills: List[str] = []
    s, e = get_section_range(lines, "skills")
    if s != -1:
        raw = " ".join(lines[s + 1:e])
        # Split by bullets / pipes / commas / slashes
        parts = re.split(r"[,|•·/\n;]", raw)
        for p in parts:
            p = p.strip(" :-–—")
            if 1 < len(p) < 40 and not p.endswith(":"):
                skills.append(p)
    # Deduplicate + also fuzzy-match common skills from full text
    text_low = text.lower()
    for cs in COMMON_SKILLS:
        if cs in text_low and not any(s.lower() == cs for s in skills):
            skills.append(cs.title() if " " not in cs else cs.title())
    # Dedupe, keep order
    seen = set(); out = []
    for s_ in skills:
        k = s_.lower().strip()
        if k and k not in seen:
            seen.add(k); out.append(s_)
    return out[:30]


def parse_experience(lines: List[str]) -> List[Dict[str, Any]]:
    s, e = get_section_range(lines, "experience")
    if s == -1:
        return []
    sec = lines[s + 1:e]
    entries: List[Dict[str, Any]] = []
    current: Dict[str, Any] = {}
    date_re = re.compile(r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|20\d{2}|19\d{2}|present|current)", re.I)

    for ln in sec:
        bullet = False
        if ln.startswith(("•", "●", "·", "-", "*", "–")) or re.match(r"^\d+\.\s", ln):
            bullet = True
        if bullet:
            bline = ln.lstrip("•●·*-–—0123456789. ").strip()
            if current:
                current.setdefault("bullets", []).append(bline)
        elif date_re.search(ln) and len(ln) < 80:
            # Likely a role/company/date header
            if current:
                entries.append(current)
            current = {"duration": ln.strip()}
            # Try previous line as role if we haven't set it
            if entries and not current.get("role"):
                pass
        elif len(ln) < 120 and not ln.endswith("."):
            # Role / company header
            if current and "role" in current and "company" in current:
                entries.append(current)
                current = {}
            if "role" not in current:
                current["role"] = ln.strip()
            elif "company" not in current:
                current["company"] = ln.strip()
        else:
            # Description line
            if current:
                current["description"] = (current.get("description") or "") + " " + ln
    if current:
        entries.append(current)
    return entries[:8]


def parse_education(lines: List[str]) -> List[Dict[str, Any]]:
    s, e = get_section_range(lines, "education")
    if s == -1:
        return []
    sec = lines[s + 1:e]
    entries: List[Dict[str, Any]] = []
    current: Dict[str, Any] = {}
    date_re = re.compile(r"20\d{2}|19\d{2}|present", re.I)
    for ln in sec:
        if date_re.search(ln) and len(ln) < 120:
            if current:
                entries.append(current)
                current = {}
            current["duration"] = ln.strip()
        elif "," in ln and len(ln) < 150:
            if "school" not in current:
                current["school"] = ln.strip()
            elif "degree" not in current:
                current["degree"] = ln.strip()
        elif len(ln) < 150:
            if "degree" not in current:
                current["degree"] = ln.strip()
            elif "school" not in current:
                current["school"] = ln.strip()
    if current:
        entries.append(current)
    return entries[:5]


def parse_resume_bytes(data: bytes, filename: str) -> Dict[str, Any]:
    filename_low = (filename or "").lower()
    if filename_low.endswith(".pdf"):
        text = extract_text_from_pdf(data)
    elif filename_low.endswith(".docx"):
        text = extract_text_from_docx(data)
    else:
        raise ValueError("Only .pdf and .docx supported")

    if not text.strip():
        raise ValueError("No extractable text found")

    lines = split_lines(text)
    personal = parse_personal(text, lines)
    summary = parse_summary(lines)
    skills = parse_skills(text, lines)
    experience = parse_experience(lines)
    education = parse_education(lines)

    confidence = 0
    if personal.get("email"):   confidence += 20
    if personal.get("phone"):   confidence += 15
    if personal.get("fullName"):confidence += 15
    if summary:                 confidence += 10
    if skills:                  confidence += 15
    if experience:              confidence += 15
    if education:               confidence += 10

    return {
        "confidence": confidence,
        "raw_text_preview": text[:500],
        "data": {
            "personal": personal,
            "summary": summary,
            "skills": skills,
            "experience": experience,
            "education": education,
            "projects": [],
            "certifications": [],
            "languages": [],
            "hobbies": [],
        },
    }
