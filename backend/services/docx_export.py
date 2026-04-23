"""DOCX resume export using python-docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from typing import Dict


def _rgb(hex_color: str):
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


ACCENT = "#FF4400"
DARK = "#0D0D12"


def _add_heading(doc: Document, text: str, level: int = 2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = _rgb(ACCENT)
    # Bottom border via paragraph style workaround
    return p


def _add_rule(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 60)
    run.font.size = Pt(6)
    run.font.color.rgb = _rgb("#E5E7EB")


def build_resume_docx(resume_doc: Dict) -> bytes:
    d = resume_doc.get("data") or {}
    p_info = d.get("personal") or {}

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # ----- HEADER -----
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(p_info.get("fullName") or "Your Name")
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = _rgb(DARK)

    contact_parts = [x for x in [
        p_info.get("email"), p_info.get("phone"),
        p_info.get("location"), p_info.get("linkedin"),
        p_info.get("website"),
    ] if x]
    if contact_parts:
        contact_para = doc.add_paragraph(" | ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in contact_para.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = _rgb("#4A4D57")

    # ----- SECTIONS -----
    order = resume_doc.get("section_order") or [
        "summary", "experience", "skills", "education",
        "projects", "certifications", "languages",
    ]

    for sec in order:
        if sec == "summary" and d.get("summary"):
            _add_heading(doc, "Professional Summary")
            _add_rule(doc)
            doc.add_paragraph(d["summary"])

        elif sec == "experience" and d.get("experience"):
            _add_heading(doc, "Work Experience")
            _add_rule(doc)
            for exp in d["experience"]:
                role_para = doc.add_paragraph()
                role_para.paragraph_format.space_before = Pt(6)
                r1 = role_para.add_run(exp.get("role") or "")
                r1.bold = True
                r1.font.size = Pt(11)
                r1.font.color.rgb = _rgb(DARK)
                if exp.get("company"):
                    r2 = role_para.add_run(f"  —  {exp['company']}")
                    r2.font.size = Pt(11)
                if exp.get("duration"):
                    dur_para = doc.add_paragraph(exp["duration"])
                    for run in dur_para.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = _rgb("#6B7280")
                if exp.get("description"):
                    doc.add_paragraph(exp["description"])
                for bullet in exp.get("bullets") or []:
                    if bullet.strip():
                        bp = doc.add_paragraph(style="List Bullet")
                        bp.add_run(bullet)

        elif sec == "education" and d.get("education"):
            _add_heading(doc, "Education")
            _add_rule(doc)
            for edu in d["education"]:
                ep = doc.add_paragraph()
                ep.paragraph_format.space_before = Pt(4)
                r1 = ep.add_run(edu.get("degree") or "")
                r1.bold = True
                if edu.get("field"):
                    ep.add_run(f" in {edu['field']}")
                if edu.get("school"):
                    ep.add_run(f"  —  {edu['school']}")
                if edu.get("duration"):
                    dp = doc.add_paragraph(edu["duration"])
                    for run in dp.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = _rgb("#6B7280")

        elif sec == "skills" and d.get("skills"):
            _add_heading(doc, "Skills")
            _add_rule(doc)
            doc.add_paragraph(", ".join(d["skills"]))

        elif sec == "projects" and d.get("projects"):
            _add_heading(doc, "Projects")
            _add_rule(doc)
            for proj in d["projects"]:
                pp = doc.add_paragraph()
                pp.paragraph_format.space_before = Pt(4)
                r1 = pp.add_run(proj.get("name") or "")
                r1.bold = True
                r1.font.color.rgb = _rgb(ACCENT)
                if proj.get("description"):
                    doc.add_paragraph(proj["description"])
                for bullet in proj.get("bullets") or []:
                    if bullet.strip():
                        bp = doc.add_paragraph(style="List Bullet")
                        bp.add_run(bullet)

        elif sec == "certifications" and d.get("certifications"):
            _add_heading(doc, "Certifications")
            _add_rule(doc)
            for cert in d["certifications"]:
                name = cert.get("name") if isinstance(cert, dict) else str(cert)
                issuer = cert.get("issuer") if isinstance(cert, dict) else ""
                year = cert.get("year") if isinstance(cert, dict) else ""
                parts = [x for x in [name, issuer, str(year) if year else ""] if x]
                doc.add_paragraph("  •  " + "  ·  ".join(parts))

        elif sec == "languages" and d.get("languages"):
            _add_heading(doc, "Languages")
            _add_rule(doc)
            langs = d["languages"]
            if isinstance(langs, list) and all(isinstance(l, dict) for l in langs):
                doc.add_paragraph(", ".join(f"{l.get('language', '')} ({l.get('proficiency', '')})" for l in langs))
            else:
                doc.add_paragraph(", ".join(str(l) for l in langs))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
